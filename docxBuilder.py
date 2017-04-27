#coding=utf-8
import os
import csv
import json
from docx import Document
#from docx.shared import Pt, Inches, RGBColor
#from docx.oxml.ns import qn
#from collections import OrderedDict

"""
根据 coms_info_complete.json 文件生成 docx 格式 word 账单
"""

# path
#cur_dir = os.path.dirname(os.path.abspath(__file__))
#pic_dir = os.path.join(cur_dir, 'picture')


class DocxBuilder:

    def __init__(self, com_name):
        self.com_name = com_name
        print u'Payment Order {com_name} Init ...'.format(com_name=self.com_name)
        self.jsonin = 'coms_info_complete.json'
        self.docxout = u'{com_name}_付款通知.docx'.format(com_name=self.com_name)
        self.doc = Document('template.docx')
        
    def getComInfoByJson(self):
        with open(self.jsonin, 'rb') as f:
            data = json.load(f)
        com_info = data['com_info']
        return com_info
    
    def setParaStyle(self, para, item):
        com_info = self.getComInfoByJson()
        com_info = com_info[self.com_name]
        paraStyle = com_info['style']
        para.style = paraStyle.get(item) or u'默认微软雅黑'

    def printComInfo(self):
        com_info = self.getComInfoByJson()
        com_info = com_info[self.com_name]
        for item in com_info['order']:
            para = self.doc.add_paragraph()
            self.setParaStyle(para, item)
            #DONE: setParaStyle()
            run = para.add_run(u'{key}：{val}'.format(key=item, val=com_info[item]))
        #debug
        #doc.save('docxOut.docx')
        #exit()
    #printComInfo('coms_info.json', u'斯雷康', doc)

    def getPayOverviewTab(self):
        with open(self.jsonin, 'rb') as f:
            data = json.load(f)
        pay_overview_tab = data['common']['tables']['payment_overview_table']
        return pay_overview_tab
        
    
    def printPayOverviewTab(self):
        pay_overview_tab = self.getPayOverviewTab()
        #TODO: set Data into pay_overview_tab
        table = self.doc.add_table(rows=1, cols=2)
        table.style = pay_overview_tab.get('style').get('tableStyle')
        #DONE: setStyle(table)
        header_cell = table.cell(0, 0).merge(table.cell(0, 1))
        header_cell.text = pay_overview_tab['header']
        # only one item 
        for item in pay_overview_tab['items']:
            for col in pay_overview_tab['order']:
                cells = table.add_row().cells
                cells[0].text = col
                cells[1].text = item[col]
        #debug
        #doc.save('docxOut.docx')
        #exit()
    #printPayOverviewTab('coms_info.json', doc)

    def getAboveBillPara(self):
        with open(self.jsonin, 'rb') as f:
            data = json.load(f)
        above_bill_para = data['common']['paragraphs']['above_billing_paragraph']
        return above_bill_para

    def printAboveBillPara(self):
        above_bill_para = self.getAboveBillPara()
        #TODO: set Data into above_bill_para
        # one paragraph only
        para = self.doc.add_paragraph()
        para.style = above_bill_para.get('style') or u'默认微软雅黑'
        #DONE: setParaStyle()
        for item in above_bill_para['order']:
            run = para.add_run(u'{key}：{val}'.format(key=item, val=above_bill_para[item]))
        #debug
        #doc.save('docxOut.docx')
        #exit()
    #printAboveBillPara('coms_info.json', doc)

    def getBillOverviewTab(self):
        with open(self.jsonin, 'rb') as f:
            data = json.load(f)
        bill_overview_tab = data['common']['tables']['billing_overview_table']
        return bill_overview_tab

    def printBillOverviewTab(self):
        bill_overview_tab = self.getBillOverviewTab()
        #TODO: set Data into bill_overview_tab
        table = self.doc.add_table(rows=1, cols=2)
        table.style = bill_overview_tab.get('style').get('tableStyle')
        #DONE: setStyle(table)
        header_cell = table.cell(0, 0).merge(table.cell(0, 1))
        header_cell.text = bill_overview_tab['header']
        # only one item 
        for item in bill_overview_tab['items']:
            for col in bill_overview_tab['order']:
                cells = table.add_row().cells
                cells[0].text = col
                cells[1].text = item[col]
        #debug
        #doc.save('docxOut.docx')
        #exit()
    #printBillOverviewTab('coms_info.json', doc)

    def getBillDetailTab(self):
        with open(self.jsonin, 'rb') as f:
            data = json.load(f)
        bill_detail_tab = data['common']['tables']['billing_detail_table']
        return bill_detail_tab

    def printBillDetailTab(self):
        bill_detail_tab = self.getBillDetailTab()
        #TODO: set Data into bill_detail_tab
        table = self.doc.add_table(rows=1, cols=2)
        table.style = bill_detail_tab.get('style').get('tableStyle')
        #DONE: setStyle(table)
        header_cell = table.cell(0, 0).merge(table.cell(0, 1))
        header_cell.text = bill_detail_tab['header']
        # 按 productCode 遍历付款项
        for item in bill_detail_tab['items']:
            # print in this order
            for col in item['order']:
                cells = table.add_row().cells
                cells[0].text = col
                #debug print col
                cells[1].text = item[col]
        #debug
        #doc.save('docxOut.docx')
        #exit()
    #printBillDetailTab('coms_info.json', doc)

    def getFooterParas(self):
        with open(self.jsonin, 'rb') as f:
            data = json.load(f)
        fooder_para = data['common']['paragraphs']['fooder_paragraph']
        return fooder_para

    def printFooterParas(self):
        fooder_para = self.getFooterParas()
        for item in fooder_para['order']:
            para = self.doc.add_paragraph()
            para.style = fooder_para.get('style') or u'默认微软雅黑'
            #DONE: setParaStyle
            run = para.add_run(u'{val}'.format(val=fooder_para[item]))
        #debug
        #doc.save('docxOut.docx')
        #exit()
    #printFooterParas('coms_info.json', doc)
    
    
    # 'coms_info_complete.json' 已经完成了它的使命，所以删掉
    def clear(self):
        if self.jsonin == 'coms_info_complete.json':
            os.remove(self.jsonin)
            print '{csvin} cleared!'.format(csvin=self.jsonin)
    
    
    def run(self):
        self.printComInfo()
        self.doc.add_paragraph()
        self.printPayOverviewTab()
        self.doc.add_page_break()
        self.doc.add_paragraph()
        self.printAboveBillPara()
        self.printBillOverviewTab()
        self.doc.add_paragraph()
        self.printBillDetailTab()
        self.doc.add_page_break()
        self.printFooterParas()
        self.doc.save(self.docxout)
        print u'Payment Order {com_name} Create Complete！'.format(com_name=self.com_name)
        #self.clear()
        print ''


if __name__ == '__main__':
    com_name = u'斯雷康'
    docxBuilder = DocxBuilder(com_name)
    docxBuilder.run()
